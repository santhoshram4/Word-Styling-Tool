import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def remove_table_borders(table):
    tbl = table._element
    tblPr = tbl.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil') 
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_column_widths_22_78(table):
    widths = [Inches(1.43), Inches(5.07)] 
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def clear_all_footers(doc):
    for section in doc.sections:
        footer = section.footer
        for param in footer.paragraphs:
            p_element = param._element
            p_element.getparent().remove(p_element)

def add_styled_page_break(paragraph, page_num, citation_name):
    """
    Input-la page break aagura idathula grey background oda 'Page X of [citation]' insert pannum.
    """
    new_p = paragraph.insert_paragraph_before()
    
    # Grey Shading (Background color)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9D9D9') 
    new_p._element.get_or_add_pPr().append(shd)

    # Paragraph properties for spacing
    new_p.paragraph_format.space_before = Pt(6)
    new_p.paragraph_format.space_after = Pt(6)

    run = new_p.add_run(f"Page {page_num} of [{citation_name}]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 128) # Blue text

def apply_legal_template_final(doc):
    clear_all_footers(doc)
    
    paragraphs = list(doc.paragraphs)
    anchor_index = -1
    
    # 1. Header/Front page cleaning
    for i, para in enumerate(paragraphs):
        clean_text = para.text.replace(" ", "").replace("-", "").upper()
        if any(kw in clean_text for kw in ["JUDGMENT", "INTRODUCTION", "BACKGROUND"]):
            anchor_index = i
            break
            
    if anchor_index != -1:
        for i in range(anchor_index):
            p = paragraphs[i]._element
            p.getparent().remove(p)

    # 2. Add New Header
    header_para = doc.paragraphs[0].insert_paragraph_before()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    case_name = "Name of the case"
    citation_text = "citation_name" 

    run1 = header_para.add_run(case_name)
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(14) 
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(0, 0, 128)
    
    header_para.add_run().add_break() 
    
    run2 = header_para.add_run(citation_text)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(14)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0, 0, 128)

    # 3. Add Table
    table_obj = doc.add_table(rows=15, cols=2)
    remove_table_borders(table_obj)
    set_column_widths_22_78(table_obj) 
    header_para._element.addnext(table_obj._element)

    # 4. Page Numbering Logic (Detection focus)
    current_page = 1
    judgment_found = False
    
    for p in doc.paragraphs:
        # Check if this paragraph contains a page break marker from the input
        # w:lastRenderedPageBreak is used when Word auto-calculates pages
        has_break = False
        if '\f' in p.text: # Manual break
            has_break = True
        elif p._element.xpath('.//w:br[@w:type="page"]'): # Explicit page break
            has_break = True
        elif p._element.xpath('.//w:lastRenderedPageBreak'): # Auto-generated break from Word
            has_break = True

        # First Page: Table-ku kila Judgment start-la
        if not judgment_found and "JUDGMENT" in p.text.upper():
            add_styled_page_break(p, current_page, citation_text)
            current_page += 1
            judgment_found = True
            continue
        
        # Following Pages: Input-la enga break iruko anga
        if has_break:
            add_styled_page_break(p, current_page, citation_text)
            current_page += 1

    # Table Labels setup
    labels = ["Reported in:", "Case No:", "Judgment Date(s):", "Hearing Date(s):",
              "Marked as:", "Country:", "Jurisdiction:", "Division:", "Judge:",
              "Bench:", "Parties:", "Appearance:", "Categories:", "Function:", "Relevant Legislation:"]
    
    for r, label in enumerate(labels):
        cell_left = table_obj.cell(r, 0)
        p_l = cell_left.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(4) 
        rl = p_l.add_run(label)
        rl.font.bold = True
        rl.font.name = 'Times New Roman'
        rl.font.size = Pt(11) 
        
        cell_right = table_obj.cell(r, 1)
        p_r = cell_right.paragraphs[0]
        p_r.paragraph_format.space_after = Pt(4)
        val = ""
        if label == "Reported in:": val = "Kenya Judgments Online, a LexisNexis Electronic Law Report Series"
        elif label == "Marked as:": val = "Unmarked"
        elif label == "Country:": val = "Kenya"
        elif label in ["Categories:", "Function:", "Relevant Legislation:"]: val = "NA"
        
        rr = p_r.add_run(val)
        rr.font.name = 'Times New Roman'
        rr.font.size = Pt(11)

def apply_checklist_rules(doc):
    for para in doc.paragraphs:
        if any(run.font.color.rgb == RGBColor(0, 0, 128) for run in para.runs):
            continue
        for run in para.runs:
            u, b, it = run.underline, run.bold, run.italic
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.underline, run.bold, run.italic = u, b, it

    replacements = {
        r'\bfootnote\b': 'fn', r'\bsection\b': 's', r'\bparagraph\b': 'para',
        r'\bregulation\b': 'reg', r'\bsubsection\b': 'ss', r'\brule\b': 'r',
        r' percent': '%', r'(\d+)\s%': r'\1%', r'(\d+),00': r'\1'
    }

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text.replace("  ", " ")
            for pattern, replacement in replacements.items():
                text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
            text = re.sub(r'(\d+)\s?centimetres', r'\1cm', text)
            text = re.sub(r'(\d+)\s?kilograms', r'\1kg', text)
            text = re.sub(r'(\d+)\s?kilometers', r'\1km', text)
            text = re.sub(r'(\d+)(st|nd|rd|th)\s([A-Z][a-z]+)\s(\d{4})', r'\1 \3 \4', text)
            run.text = text

def process_batch():
    input_dir = 'input'
    output_dir = 'output'
    if not os.path.exists(output_dir): os.makedirs(output_dir)

    for filename in os.listdir(input_dir):
        if filename.lower().endswith('.docx') and not filename.startswith('~$'):
            print(f"Processing: {filename}")
            try:
                doc = Document(os.path.join(input_dir, filename))
                apply_legal_template_final(doc)
                apply_checklist_rules(doc)
                doc.save(os.path.join(output_dir, filename))
            except Exception as e:
                print(f"Error processing {filename}: {e}")

if __name__ == "__main__":
    process_batch()